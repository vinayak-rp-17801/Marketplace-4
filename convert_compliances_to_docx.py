from __future__ import annotations

import difflib
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches


SOURCE_DIR = Path("compliances")
OUTPUT_DIR = Path("compliances-docx")
REGIONAL_ZIP = Path("Regional  Compliance.zip")
SCREENSHOT_ZIPS = [Path(f"{index}.zip") for index in range(1, 6)]

DOC_CODE_OVERRIDES = {
    "b3s-krankenhaus-and-sgb-v-section-75c-compliance-for-log360-cloud": "B3S / Section 75c",
    "banca-d-italia-circular-285-2013-compliance-for-log360-cloud": "Circ. 285/2013",
    "banco-de-espana-circular-2-2016-compliance-for-log360-cloud": "BdE Circ. 2/2016",
    "desc-information-security-regulation-compliance-for-log360-cloud": "DESC ISR",
    "finma-circular-2023-1-compliance-for-log360-cloud": "FINMA 2023/1",
    "french-ciip-law-2013-compliance-for-log360-cloud": "LPM/CIIP",
    "french-sren-law-2024-compliance-for-log360-cloud": "SREN 2024",
    "ivass-regulation-38-2018-compliance-for-log360-cloud": "IVASS 38/2018",
}


def parse_markdown(content: str) -> tuple[str, str, str, list[str]]:
    lines = [line.rstrip() for line in content.splitlines()]

    title = ""
    tagline = ""
    overview_lines: list[str] = []
    key_features: list[str] = []

    section = None
    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("# ") and not title:
            title = line[2:].strip()
            section = None
            continue

        if line.startswith("**Tagline:**"):
            tagline = line.replace("**Tagline:**", "", 1).strip()
            section = None
            continue

        if line == "## Overview":
            section = "overview"
            continue

        if line == "## Key Features":
            section = "key_features"
            continue

        if section == "overview":
            overview_lines.append(line)
        elif section == "key_features" and line.startswith("- "):
            key_features.append(line[2:].strip())

    return title, tagline, "\n".join(overview_lines).strip(), key_features


def normalize(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", text.lower())


def parse_compliance_codes(regional_extract_dir: Path) -> list[tuple[str, str]]:
    source_file = regional_extract_dir / "Regional  Compliance" / "compliance_logo_sources.txt"
    entries: list[tuple[str, str]] = []
    pattern = re.compile(r"^\[[^\]]+\]\s+(.+?)\s+—\s+(.+)$")

    for line in source_file.read_text(encoding="utf-8", errors="ignore").splitlines():
        match = pattern.match(line.strip())
        if match:
            entries.append((match.group(1).strip(), match.group(2).strip()))

    return entries


def find_asset_folders(base_dir: Path) -> list[str]:
    folders: list[str] = []
    for folder in sorted(base_dir.iterdir()):
        if folder.is_dir() and not folder.name.startswith("__MACOSX"):
            folders.append(folder.name)
    return folders


def pick_best_folder(code: str, folders: list[str]) -> str:
    code_key = normalize(code)
    scored = sorted(
        (
            difflib.SequenceMatcher(None, code_key, normalize(folder)).ratio(),
            folder,
        )
        for folder in folders
    )
    return scored[-1][1]


def pick_code_for_doc(stem: str, title: str, code_entries: list[tuple[str, str]]) -> str:
    if stem in DOC_CODE_OVERRIDES:
        return DOC_CODE_OVERRIDES[stem]

    title_key = normalize(title)
    query_key = normalize(f"{stem} {title.replace(' Compliance for Log360 Cloud', '')}")

    best_score = -1.0
    best_code = ""

    for code, code_title in code_entries:
        code_key = normalize(code)
        code_title_key = normalize(code_title)

        score = max(
            difflib.SequenceMatcher(None, title_key, code_title_key).ratio(),
            difflib.SequenceMatcher(None, query_key, code_title_key).ratio(),
            difflib.SequenceMatcher(None, query_key, code_key).ratio(),
        )

        if code_key and code_key in query_key:
            score = 1.0
        elif code_title_key and (code_title_key in query_key or query_key in code_title_key):
            score = max(score, 0.95)

        if score > best_score:
            best_score = score
            best_code = code

    return best_code


def select_image(asset_dir: Path, marker: str) -> Path:
    candidates = sorted(
        path
        for path in asset_dir.iterdir()
        if path.is_file() and marker in path.stem and path.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"}
    )
    if not candidates:
        raise FileNotFoundError(f"No image found in {asset_dir} matching marker '{marker}'")
    return candidates[0]


def get_screenshots(asset_dir: Path) -> list[Path]:
    screenshots = sorted(
        path
        for path in asset_dir.iterdir()
        if path.is_file() and path.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"}
    )
    if not screenshots:
        raise FileNotFoundError(f"No screenshots found in {asset_dir}")
    return screenshots


def create_docx_from_markdown(
    markdown_path: Path,
    output_path: Path,
    logo_image: Path,
    thumbnail_image: Path,
    screenshots: list[Path],
) -> None:
    title, tagline, overview, key_features = parse_markdown(markdown_path.read_text(encoding="utf-8"))

    document = Document()

    if title:
        document.add_heading(title, level=1)

    if tagline:
        tagline_paragraph = document.add_paragraph()
        label_run = tagline_paragraph.add_run("Tagline:")
        label_run.bold = True
        tagline_paragraph.add_run(f" {tagline}")

    document.add_heading("Overview", level=2)
    if overview:
        for paragraph_text in overview.split("\n"):
            if paragraph_text.strip():
                document.add_paragraph(paragraph_text)

    if key_features:
        document.add_heading("Key Features", level=2)
        for feature in key_features:
            document.add_paragraph(feature, style="List Bullet")

    document.add_picture(str(logo_image), width=Inches(180 / 96))
    document.add_picture(str(thumbnail_image), width=Inches(740 / 96))

    document.add_heading("Screenshots", level=2)
    for screenshot in screenshots:
        document.add_picture(str(screenshot), width=Inches(6.5))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory() as temp_dir_str:
        temp_dir = Path(temp_dir_str)

        regional_extract_dir = temp_dir / "regional"
        with zipfile.ZipFile(REGIONAL_ZIP) as regional_zip:
            regional_zip.extractall(regional_extract_dir)

        screenshot_extract_dirs: list[Path] = []
        for screenshot_zip_path in SCREENSHOT_ZIPS:
            extract_dir = temp_dir / screenshot_zip_path.stem
            with zipfile.ZipFile(screenshot_zip_path) as screenshot_zip:
                screenshot_zip.extractall(extract_dir)
            screenshot_extract_dirs.append(extract_dir)

        code_entries = parse_compliance_codes(regional_extract_dir)
        regional_asset_root = regional_extract_dir / "Regional  Compliance"
        regional_folders = find_asset_folders(regional_asset_root)

        screenshot_folders: dict[str, Path] = {}
        for screenshot_extract_dir in screenshot_extract_dirs:
            zip_folder = screenshot_extract_dir / screenshot_extract_dir.name
            for folder_name in find_asset_folders(zip_folder):
                screenshot_folders[folder_name] = zip_folder / folder_name
        screenshot_folder_names = sorted(screenshot_folders)

        for markdown_file in sorted(SOURCE_DIR.glob("*.md")):
            title, _, _, _ = parse_markdown(markdown_file.read_text(encoding="utf-8"))
            code = pick_code_for_doc(markdown_file.stem, title, code_entries)

            regional_folder = pick_best_folder(code, regional_folders)
            screenshot_folder = pick_best_folder(code, screenshot_folder_names)

            logo_image = select_image(regional_asset_root / regional_folder, "_180x180")
            thumbnail_image = select_image(regional_asset_root / regional_folder, "_740x340")
            screenshots = get_screenshots(screenshot_folders[screenshot_folder])

            output_file = OUTPUT_DIR / f"{markdown_file.stem}.docx"
            create_docx_from_markdown(
                markdown_file,
                output_file,
                logo_image=logo_image,
                thumbnail_image=thumbnail_image,
                screenshots=screenshots,
            )

    with zipfile.ZipFile("compliances-docx.zip", "w", compression=zipfile.ZIP_DEFLATED) as output_zip:
        for docx_file in sorted(OUTPUT_DIR.glob("*.docx")):
            output_zip.write(docx_file, arcname=f"compliances-docx/{docx_file.name}")


if __name__ == "__main__":
    main()
